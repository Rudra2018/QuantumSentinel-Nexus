#!/usr/bin/env python3
"""
QuantumSentinel-Nexus Report Formatters
Professional report formatting with real evidence integration
"""

import asyncio
import logging
import json
import base64
from typing import Dict, List, Any, Optional
from dataclasses import dataclass
from pathlib import Path
from datetime import datetime, timezone
import hashlib
import uuid
from io import BytesIO

# Document generation imports
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, Table, TableStyle, PageBreak
from reportlab.lib.colors import black, red, orange, yellow, green, blue, grey
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
from reportlab.graphics.shapes import Drawing, Rect
from reportlab.graphics.charts.barcharts import VerticalBarChart
from reportlab.graphics.charts.piecharts import Pie
from reportlab.lib.utils import ImageReader

# HTML/PDF generation
from jinja2 import Environment, BaseLoader, Template
from weasyprint import HTML, CSS
import markdown
from markdown.extensions import codehilite, tables, toc

# Office document generation
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import openpyxl
from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
from openpyxl.chart import BarChart, PieChart, Reference

# Image processing
from PIL import Image, ImageDraw, ImageFont
import matplotlib.pyplot as plt
import seaborn as sns

@dataclass
class ReportSection:
    """Individual report section"""
    title: str
    content: str
    section_type: str  # text, table, chart, image, code
    metadata: Dict[str, Any]
    order: int = 0

@dataclass
class ReportTemplate:
    """Report template definition"""
    template_id: str
    name: str
    description: str
    sections: List[str]
    required_data: List[str]
    output_formats: List[str]
    style_config: Dict[str, Any]

class BaseFormatter:
    """Base class for all report formatters"""

    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.{self.__class__.__name__}")

    async def format_report(self, report_data: Dict[str, Any], **kwargs) -> str:
        """Format report data into specific output format"""
        raise NotImplementedError("Subclasses must implement format_report method")

    def _prepare_data(self, report_data: Dict[str, Any]) -> Dict[str, Any]:
        """Prepare and sanitize data for formatting"""
        # Common data preparation logic
        prepared_data = report_data.copy()

        # Ensure timestamps are formatted consistently
        if 'timestamp' in prepared_data:
            if isinstance(prepared_data['timestamp'], datetime):
                prepared_data['timestamp'] = prepared_data['timestamp'].strftime('%Y-%m-%d %H:%M:%S UTC')

        # Sanitize HTML content
        if 'html_content' in prepared_data:
            prepared_data['html_content'] = self._sanitize_html(prepared_data['html_content'])

        return prepared_data

    def _sanitize_html(self, html_content: str) -> str:
        """Sanitize HTML content"""
        # Basic HTML sanitization - in production, use a proper library like bleach
        import html
        return html.escape(html_content)

    def _generate_report_id(self) -> str:
        """Generate unique report ID"""
        return f"QS-{datetime.now().strftime('%Y%m%d')}-{uuid.uuid4().hex[:8].upper()}"

class HTMLFormatter(BaseFormatter):
    """HTML report formatter with professional styling"""

    def __init__(self):
        super().__init__()
        self.template_loader = self._setup_template_loader()

    def _setup_template_loader(self):
        """Setup Jinja2 template loader"""
        template_str = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{ report_title }}</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            color: #333;
            background-color: #f8f9fa;
        }
        .container { max-width: 1200px; margin: 0 auto; padding: 20px; }
        .header {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 40px 0;
            margin-bottom: 30px;
            border-radius: 10px;
        }
        .header h1 { font-size: 2.5em; margin-bottom: 10px; text-align: center; }
        .header .subtitle { font-size: 1.2em; text-align: center; opacity: 0.9; }
        .meta-info {
            background: white;
            padding: 20px;
            border-radius: 8px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
            margin-bottom: 30px;
        }
        .meta-grid { display: grid; grid-template-columns: repeat(auto-fit, minmax(250px, 1fr)); gap: 20px; }
        .meta-item { padding: 15px; background: #f8f9fa; border-radius: 5px; }
        .meta-item strong { color: #495057; }
        .section {
            background: white;
            margin-bottom: 30px;
            border-radius: 8px;
            overflow: hidden;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }
        .section-header {
            background: #343a40;
            color: white;
            padding: 20px;
            font-size: 1.3em;
            font-weight: 600;
        }
        .section-content { padding: 30px; }
        .vulnerability {
            border-left: 4px solid #dc3545;
            padding: 20px;
            margin: 20px 0;
            background: #fff5f5;
            border-radius: 0 8px 8px 0;
        }
        .vulnerability.high { border-left-color: #fd7e14; background: #fff8f0; }
        .vulnerability.medium { border-left-color: #ffc107; background: #fffdf0; }
        .vulnerability.low { border-left-color: #28a745; background: #f0fff0; }
        .severity-badge {
            display: inline-block;
            padding: 6px 12px;
            border-radius: 20px;
            font-size: 0.9em;
            font-weight: 600;
            text-transform: uppercase;
        }
        .severity-critical { background: #dc3545; color: white; }
        .severity-high { background: #fd7e14; color: white; }
        .severity-medium { background: #ffc107; color: #856404; }
        .severity-low { background: #28a745; color: white; }
        .evidence-item {
            background: #f8f9fa;
            border: 1px solid #dee2e6;
            border-radius: 5px;
            padding: 15px;
            margin: 10px 0;
        }
        .code-block {
            background: #f8f9fa;
            border: 1px solid #e9ecef;
            border-radius: 5px;
            padding: 15px;
            font-family: 'Courier New', monospace;
            overflow-x: auto;
            white-space: pre-wrap;
        }
        .chart-container { text-align: center; margin: 20px 0; }
        .chart-container img { max-width: 100%; height: auto; }
        .recommendations {
            background: #e7f3ff;
            border: 1px solid #b8daff;
            border-radius: 5px;
            padding: 20px;
            margin: 20px 0;
        }
        .recommendations h3 { color: #004085; margin-bottom: 15px; }
        .recommendations ul { margin-left: 20px; }
        .recommendations li { margin-bottom: 8px; }
        .table-container { overflow-x: auto; margin: 20px 0; }
        table { width: 100%; border-collapse: collapse; }
        th, td { padding: 12px; text-align: left; border-bottom: 1px solid #dee2e6; }
        th { background-color: #f8f9fa; font-weight: 600; }
        .footer {
            background: #343a40;
            color: white;
            padding: 30px 0;
            text-align: center;
            margin-top: 50px;
            border-radius: 10px;
        }
        .footer p { margin-bottom: 5px; }
        .watermark {
            position: fixed;
            bottom: 20px;
            right: 20px;
            opacity: 0.7;
            font-size: 0.8em;
            color: #6c757d;
        }
        @media print {
            .container { max-width: none; margin: 0; padding: 0; }
            .header, .footer { break-inside: avoid; }
            .section { break-inside: avoid; page-break-inside: avoid; }
        }
    </style>
</head>
<body>
    <div class="container">
        <header class="header">
            <h1>{{ report_title }}</h1>
            <div class="subtitle">{{ report_subtitle }}</div>
        </header>

        <div class="meta-info">
            <div class="meta-grid">
                <div class="meta-item">
                    <strong>Report ID:</strong> {{ report_id }}
                </div>
                <div class="meta-item">
                    <strong>Generated:</strong> {{ generated_at }}
                </div>
                <div class="meta-item">
                    <strong>Target:</strong> {{ target_info.get('target', 'N/A') }}
                </div>
                <div class="meta-item">
                    <strong>Assessment Type:</strong> {{ target_info.get('target_type', 'N/A') }}
                </div>
                <div class="meta-item">
                    <strong>Total Findings:</strong> {{ vulnerabilities|length }}
                </div>
                <div class="meta-item">
                    <strong>Execution Time:</strong> {{ execution_time_formatted }}
                </div>
            </div>
        </div>

        {% if summary_stats %}
        <div class="section">
            <div class="section-header">Executive Summary</div>
            <div class="section-content">
                <div class="meta-grid">
                    <div class="meta-item">
                        <strong>Risk Score:</strong> {{ summary_stats.risk_score }}/10
                    </div>
                    <div class="meta-item">
                        <strong>Critical Findings:</strong> {{ summary_stats.severity_distribution.get('critical', 0) }}
                    </div>
                    <div class="meta-item">
                        <strong>High Severity:</strong> {{ summary_stats.severity_distribution.get('high', 0) }}
                    </div>
                    <div class="meta-item">
                        <strong>Average CVSS:</strong> {{ summary_stats.get('average_cvss_score', 'N/A') }}
                    </div>
                </div>

                {% if recommendations %}
                <div class="recommendations">
                    <h3>Key Recommendations</h3>
                    <ul>
                        {% for recommendation in recommendations %}
                        <li>{{ recommendation }}</li>
                        {% endfor %}
                    </ul>
                </div>
                {% endif %}
            </div>
        </div>
        {% endif %}

        {% if vulnerabilities %}
        <div class="section">
            <div class="section-header">Vulnerability Findings</div>
            <div class="section-content">
                {% for vuln in vulnerabilities %}
                <div class="vulnerability {{ vuln.get('severity', 'medium').lower() }}">
                    <h3>{{ vuln.get('title', 'Unknown Vulnerability') }}
                        <span class="severity-badge severity-{{ vuln.get('severity', 'medium').lower() }}">
                            {{ vuln.get('severity', 'Medium') }}
                        </span>
                    </h3>

                    <p><strong>Type:</strong> {{ vuln.get('vulnerability_type', 'Unknown') }}</p>
                    {% if vuln.get('cvss_score') %}
                    <p><strong>CVSS Score:</strong> {{ vuln.cvss_score }}/10.0</p>
                    {% endif %}

                    <h4>Description</h4>
                    <p>{{ vuln.get('description', 'No description available') }}</p>

                    {% if vuln.get('proof_of_concept') %}
                    <h4>Proof of Concept</h4>
                    <div class="code-block">{{ vuln.proof_of_concept }}</div>
                    {% endif %}

                    {% if vuln.get('impact') %}
                    <h4>Impact</h4>
                    <p>{{ vuln.impact }}</p>
                    {% endif %}

                    {% if vuln.get('remediation') %}
                    <h4>Remediation</h4>
                    <p>{{ vuln.remediation }}</p>
                    {% endif %}
                </div>
                {% endfor %}
            </div>
        </div>
        {% endif %}

        {% if evidence_summary %}
        <div class="section">
            <div class="section-header">Evidence Collection Summary</div>
            <div class="section-content">
                <div class="meta-grid">
                    <div class="meta-item">
                        <strong>Evidence Items:</strong> {{ evidence_summary.get('total_evidence_items', 0) }}
                    </div>
                    <div class="meta-item">
                        <strong>Collection ID:</strong> {{ evidence_summary.get('collection_id', 'N/A') }}
                    </div>
                    <div class="meta-item">
                        <strong>Total Size:</strong> {{ evidence_summary.get('total_size_bytes', 0) }} bytes
                    </div>
                </div>

                {% if evidence_summary.get('evidence_by_type') %}
                <h4>Evidence Types</h4>
                <div class="table-container">
                    <table>
                        <thead>
                            <tr>
                                <th>Type</th>
                                <th>Count</th>
                                <th>Total Size</th>
                            </tr>
                        </thead>
                        <tbody>
                            {% for type_name, type_data in evidence_summary.evidence_by_type.items() %}
                            <tr>
                                <td>{{ type_name.replace('_', ' ').title() }}</td>
                                <td>{{ type_data.count }}</td>
                                <td>{{ type_data.total_size }} bytes</td>
                            </tr>
                            {% endfor %}
                        </tbody>
                    </table>
                </div>
                {% endif %}
            </div>
        </div>
        {% endif %}

        {% if charts %}
        <div class="section">
            <div class="section-header">Visualizations</div>
            <div class="section-content">
                {% for chart_name, chart_data in charts.items() %}
                <div class="chart-container">
                    <h4>{{ chart_name.replace('_', ' ').title() }}</h4>
                    <img src="{{ chart_data }}" alt="{{ chart_name }}">
                </div>
                {% endfor %}
            </div>
        </div>
        {% endif %}

        <footer class="footer">
            <p>QuantumSentinel-Nexus Security Assessment Report</p>
            <p>Generated by QuantumSentinel-Nexus Ultimate Security Testing Platform</p>
            <p>Report ID: {{ report_id }} | Generated: {{ generated_at }}</p>
        </footer>
    </div>

    <div class="watermark">
        QuantumSentinel-Nexus | Confidential
    </div>
</body>
</html>
        """
        return Environment(loader=BaseLoader()).from_string(template_str)

    async def format_report(self, report_data: Dict[str, Any], template_name: str = "comprehensive_report") -> str:
        """Format report data as HTML"""
        try:
            prepared_data = self._prepare_data(report_data)

            # Generate report context
            context = {
                'report_id': self._generate_report_id(),
                'report_title': 'QuantumSentinel Security Assessment Report',
                'report_subtitle': 'Comprehensive Security Analysis',
                'generated_at': datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC'),
                'target_info': prepared_data.get('target_info', {}),
                'vulnerabilities': prepared_data.get('vulnerabilities_found', []),
                'summary_stats': prepared_data.get('summary_stats', {}),
                'evidence_summary': prepared_data.get('evidence_summary', {}),
                'recommendations': prepared_data.get('recommendations', []),
                'charts': prepared_data.get('charts', {}),
                'execution_time_formatted': self._format_execution_time(prepared_data.get('execution_time_seconds', 0))
            }

            # Render template
            html_content = self.template_loader.render(**context)

            return html_content

        except Exception as e:
            self.logger.error(f"Error formatting HTML report: {e}")
            raise

    def _format_execution_time(self, seconds: float) -> str:
        """Format execution time in human-readable format"""
        if seconds < 60:
            return f"{seconds:.1f} seconds"
        elif seconds < 3600:
            minutes = seconds / 60
            return f"{minutes:.1f} minutes"
        else:
            hours = seconds / 3600
            return f"{hours:.1f} hours"

class PDFFormatter(BaseFormatter):
    """PDF report formatter using ReportLab"""

    def __init__(self):
        super().__init__()
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()

    def _setup_custom_styles(self):
        """Setup custom PDF styles"""
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=blue
        ))

        self.styles.add(ParagraphStyle(
            name='VulnerabilityTitle',
            parent=self.styles['Heading2'],
            fontSize=14,
            spaceBefore=20,
            spaceAfter=10,
            textColor=red
        ))

        self.styles.add(ParagraphStyle(
            name='CodeBlock',
            parent=self.styles['Normal'],
            fontSize=9,
            fontName='Courier',
            leftIndent=20,
            backgroundColor=grey
        ))

    async def format_report(self, report_data: Dict[str, Any], output_path: str, template_name: str = "comprehensive_report"):
        """Format report data as PDF"""
        try:
            prepared_data = self._prepare_data(report_data)

            # Create PDF document
            doc = SimpleDocTemplate(
                output_path,
                pagesize=letter,
                topMargin=1*inch,
                bottomMargin=1*inch,
                leftMargin=1*inch,
                rightMargin=1*inch
            )

            # Build content
            story = []

            # Title page
            story.extend(self._create_title_page(prepared_data))

            # Executive summary
            story.extend(self._create_executive_summary(prepared_data))

            # Vulnerabilities section
            story.extend(self._create_vulnerabilities_section(prepared_data))

            # Evidence section
            story.extend(self._create_evidence_section(prepared_data))

            # Build PDF
            doc.build(story)

            self.logger.info(f"PDF report generated: {output_path}")

        except Exception as e:
            self.logger.error(f"Error formatting PDF report: {e}")
            raise

    def _create_title_page(self, data: Dict[str, Any]) -> List:
        """Create PDF title page"""
        story = []

        # Title
        title = Paragraph("QuantumSentinel Security Assessment Report", self.styles['CustomTitle'])
        story.append(title)
        story.append(Spacer(1, 20))

        # Metadata table
        target_info = data.get('target_info', {})
        metadata = [
            ['Report ID:', self._generate_report_id()],
            ['Generated:', datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')],
            ['Target:', target_info.get('target', 'N/A')],
            ['Assessment Type:', target_info.get('target_type', 'N/A')],
            ['Total Findings:', str(len(data.get('vulnerabilities_found', [])))],
            ['Execution Time:', f"{data.get('execution_time_seconds', 0):.1f} seconds"]
        ]

        table = Table(metadata, colWidths=[2*inch, 4*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), grey),
            ('TEXTCOLOR', (0, 0), (-1, -1), black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('BACKGROUND', (1, 0), (1, -1), 'white'),
            ('GRID', (0, 0), (-1, -1), 1, black)
        ]))

        story.append(table)
        story.append(PageBreak())

        return story

    def _create_executive_summary(self, data: Dict[str, Any]) -> List:
        """Create executive summary section"""
        story = []

        summary_stats = data.get('summary_stats', {})
        if not summary_stats:
            return story

        # Section title
        title = Paragraph("Executive Summary", self.styles['Heading1'])
        story.append(title)
        story.append(Spacer(1, 12))

        # Risk assessment
        risk_score = summary_stats.get('risk_score', 0)
        risk_text = f"Overall Risk Score: {risk_score}/10"
        story.append(Paragraph(risk_text, self.styles['Normal']))
        story.append(Spacer(1, 12))

        # Severity distribution
        severity_dist = summary_stats.get('severity_distribution', {})
        if severity_dist:
            story.append(Paragraph("Severity Distribution:", self.styles['Heading2']))

            severity_data = [['Severity', 'Count']]
            for severity, count in severity_dist.items():
                severity_data.append([severity.title(), str(count)])

            table = Table(severity_data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), black),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('BACKGROUND', (0, 1), (-1, -1), 'white'),
                ('GRID', (0, 0), (-1, -1), 1, black)
            ]))

            story.append(table)

        story.append(PageBreak())
        return story

    def _create_vulnerabilities_section(self, data: Dict[str, Any]) -> List:
        """Create vulnerabilities section"""
        story = []

        vulnerabilities = data.get('vulnerabilities_found', [])
        if not vulnerabilities:
            return story

        # Section title
        title = Paragraph("Vulnerability Findings", self.styles['Heading1'])
        story.append(title)
        story.append(Spacer(1, 12))

        for i, vuln in enumerate(vulnerabilities, 1):
            # Vulnerability title
            vuln_title = f"{i}. {vuln.get('title', 'Unknown Vulnerability')} ({vuln.get('severity', 'Medium')})"
            story.append(Paragraph(vuln_title, self.styles['VulnerabilityTitle']))

            # Description
            if vuln.get('description'):
                story.append(Paragraph(f"<b>Description:</b> {vuln['description']}", self.styles['Normal']))
                story.append(Spacer(1, 6))

            # CVSS Score
            if vuln.get('cvss_score'):
                story.append(Paragraph(f"<b>CVSS Score:</b> {vuln['cvss_score']}/10.0", self.styles['Normal']))
                story.append(Spacer(1, 6))

            # Proof of Concept
            if vuln.get('proof_of_concept'):
                story.append(Paragraph("<b>Proof of Concept:</b>", self.styles['Normal']))
                story.append(Paragraph(vuln['proof_of_concept'], self.styles['CodeBlock']))
                story.append(Spacer(1, 6))

            # Impact
            if vuln.get('impact'):
                story.append(Paragraph(f"<b>Impact:</b> {vuln['impact']}", self.styles['Normal']))
                story.append(Spacer(1, 6))

            # Remediation
            if vuln.get('remediation'):
                story.append(Paragraph(f"<b>Remediation:</b> {vuln['remediation']}", self.styles['Normal']))

            story.append(Spacer(1, 20))

        return story

    def _create_evidence_section(self, data: Dict[str, Any]) -> List:
        """Create evidence section"""
        story = []

        evidence_summary = data.get('evidence_summary', {})
        if not evidence_summary:
            return story

        # Section title
        title = Paragraph("Evidence Collection Summary", self.styles['Heading1'])
        story.append(title)
        story.append(Spacer(1, 12))

        # Evidence statistics
        total_items = evidence_summary.get('total_evidence_items', 0)
        collection_id = evidence_summary.get('collection_id', 'N/A')

        story.append(Paragraph(f"<b>Total Evidence Items:</b> {total_items}", self.styles['Normal']))
        story.append(Paragraph(f"<b>Collection ID:</b> {collection_id}", self.styles['Normal']))
        story.append(Spacer(1, 12))

        # Evidence by type
        evidence_by_type = evidence_summary.get('evidence_by_type', {})
        if evidence_by_type:
            story.append(Paragraph("Evidence Types:", self.styles['Heading2']))

            type_data = [['Type', 'Count', 'Total Size (bytes)']]
            for type_name, type_info in evidence_by_type.items():
                type_data.append([
                    type_name.replace('_', ' ').title(),
                    str(type_info.get('count', 0)),
                    str(type_info.get('total_size', 0))
                ])

            table = Table(type_data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), black),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('BACKGROUND', (0, 1), (-1, -1), 'white'),
                ('GRID', (0, 0), (-1, -1), 1, black)
            ]))

            story.append(table)

        return story

class JSONFormatter(BaseFormatter):
    """JSON report formatter"""

    async def format_report(self, report_data: Dict[str, Any], **kwargs) -> str:
        """Format report data as JSON"""
        try:
            prepared_data = self._prepare_data(report_data)

            # Add metadata
            formatted_report = {
                'report_metadata': {
                    'report_id': self._generate_report_id(),
                    'generated_at': datetime.now(timezone.utc).isoformat(),
                    'format': 'json',
                    'version': '1.0'
                },
                'assessment_data': prepared_data
            }

            return json.dumps(formatted_report, indent=2, default=str)

        except Exception as e:
            self.logger.error(f"Error formatting JSON report: {e}")
            raise

class XMLFormatter(BaseFormatter):
    """XML report formatter"""

    async def format_report(self, report_data: Dict[str, Any], **kwargs) -> str:
        """Format report data as XML"""
        try:
            prepared_data = self._prepare_data(report_data)

            xml_content = f"""<?xml version="1.0" encoding="UTF-8"?>
<security_assessment_report>
    <metadata>
        <report_id>{self._generate_report_id()}</report_id>
        <generated_at>{datetime.now(timezone.utc).isoformat()}</generated_at>
        <format>xml</format>
        <version>1.0</version>
    </metadata>
    <target_info>
        <target>{prepared_data.get('target_info', {}).get('target', 'N/A')}</target>
        <target_type>{prepared_data.get('target_info', {}).get('target_type', 'N/A')}</target_type>
    </target_info>
    <vulnerabilities>
"""

            vulnerabilities = prepared_data.get('vulnerabilities_found', [])
            for vuln in vulnerabilities:
                xml_content += f"""        <vulnerability>
            <title>{self._escape_xml(vuln.get('title', 'Unknown'))}</title>
            <severity>{vuln.get('severity', 'Medium')}</severity>
            <type>{vuln.get('vulnerability_type', 'Unknown')}</type>
            <cvss_score>{vuln.get('cvss_score', 'N/A')}</cvss_score>
            <description>{self._escape_xml(vuln.get('description', ''))}</description>
        </vulnerability>
"""

            xml_content += """    </vulnerabilities>
</security_assessment_report>"""

            return xml_content

        except Exception as e:
            self.logger.error(f"Error formatting XML report: {e}")
            raise

    def _escape_xml(self, text: str) -> str:
        """Escape XML special characters"""
        if not isinstance(text, str):
            text = str(text)
        return (text.replace('&', '&amp;')
                   .replace('<', '&lt;')
                   .replace('>', '&gt;')
                   .replace('"', '&quot;')
                   .replace("'", '&apos;'))

class DOCXFormatter(BaseFormatter):
    """Microsoft Word DOCX formatter"""

    async def format_report(self, report_data: Dict[str, Any], output_path: str, **kwargs) -> str:
        """Format report data as DOCX"""
        try:
            prepared_data = self._prepare_data(report_data)

            # Create document
            doc = Document()

            # Add title
            title = doc.add_heading('QuantumSentinel Security Assessment Report', level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add metadata
            target_info = prepared_data.get('target_info', {})
            doc.add_paragraph(f"Report ID: {self._generate_report_id()}")
            doc.add_paragraph(f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}")
            doc.add_paragraph(f"Target: {target_info.get('target', 'N/A')}")
            doc.add_paragraph(f"Assessment Type: {target_info.get('target_type', 'N/A')}")

            # Add vulnerabilities section
            vulnerabilities = prepared_data.get('vulnerabilities_found', [])
            if vulnerabilities:
                doc.add_heading('Vulnerability Findings', level=2)

                for i, vuln in enumerate(vulnerabilities, 1):
                    vuln_heading = doc.add_heading(f"{i}. {vuln.get('title', 'Unknown Vulnerability')}", level=3)

                    doc.add_paragraph(f"Severity: {vuln.get('severity', 'Medium')}")
                    doc.add_paragraph(f"Type: {vuln.get('vulnerability_type', 'Unknown')}")

                    if vuln.get('cvss_score'):
                        doc.add_paragraph(f"CVSS Score: {vuln['cvss_score']}/10.0")

                    if vuln.get('description'):
                        desc_p = doc.add_paragraph()
                        desc_p.add_run('Description: ').bold = True
                        desc_p.add_run(vuln['description'])

                    if vuln.get('remediation'):
                        rem_p = doc.add_paragraph()
                        rem_p.add_run('Remediation: ').bold = True
                        rem_p.add_run(vuln['remediation'])

            # Save document
            doc.save(output_path)

            return output_path

        except Exception as e:
            self.logger.error(f"Error formatting DOCX report: {e}")
            raise

class ExcelFormatter(BaseFormatter):
    """Microsoft Excel XLSX formatter"""

    async def format_report(self, report_data: Dict[str, Any], output_path: str, **kwargs) -> str:
        """Format report data as Excel"""
        try:
            prepared_data = self._prepare_data(report_data)

            # Create workbook
            wb = openpyxl.Workbook()

            # Summary sheet
            ws_summary = wb.active
            ws_summary.title = "Summary"

            # Add headers
            ws_summary['A1'] = "QuantumSentinel Security Assessment Report"
            ws_summary['A1'].font = Font(size=16, bold=True)

            # Add metadata
            ws_summary['A3'] = "Report ID:"
            ws_summary['B3'] = self._generate_report_id()
            ws_summary['A4'] = "Generated:"
            ws_summary['B4'] = datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')

            target_info = prepared_data.get('target_info', {})
            ws_summary['A5'] = "Target:"
            ws_summary['B5'] = target_info.get('target', 'N/A')
            ws_summary['A6'] = "Assessment Type:"
            ws_summary['B6'] = target_info.get('target_type', 'N/A')

            # Vulnerabilities sheet
            vulnerabilities = prepared_data.get('vulnerabilities_found', [])
            if vulnerabilities:
                ws_vulns = wb.create_sheet("Vulnerabilities")

                # Headers
                headers = ['Title', 'Severity', 'Type', 'CVSS Score', 'Description', 'Remediation']
                for col, header in enumerate(headers, 1):
                    cell = ws_vulns.cell(row=1, column=col)
                    cell.value = header
                    cell.font = Font(bold=True)
                    cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")

                # Data
                for row, vuln in enumerate(vulnerabilities, 2):
                    ws_vulns.cell(row=row, column=1, value=vuln.get('title', 'Unknown'))
                    ws_vulns.cell(row=row, column=2, value=vuln.get('severity', 'Medium'))
                    ws_vulns.cell(row=row, column=3, value=vuln.get('vulnerability_type', 'Unknown'))
                    ws_vulns.cell(row=row, column=4, value=vuln.get('cvss_score', 'N/A'))
                    ws_vulns.cell(row=row, column=5, value=vuln.get('description', ''))
                    ws_vulns.cell(row=row, column=6, value=vuln.get('remediation', ''))

                # Auto-adjust column widths
                for col in ws_vulns.columns:
                    max_length = 0
                    column = col[0].column_letter
                    for cell in col:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 50)
                    ws_vulns.column_dimensions[column].width = adjusted_width

            # Save workbook
            wb.save(output_path)

            return output_path

        except Exception as e:
            self.logger.error(f"Error formatting Excel report: {e}")
            raise

# Export main classes
__all__ = [
    'HTMLFormatter',
    'PDFFormatter',
    'JSONFormatter',
    'XMLFormatter',
    'DOCXFormatter',
    'ExcelFormatter',
    'ReportSection',
    'ReportTemplate'
]